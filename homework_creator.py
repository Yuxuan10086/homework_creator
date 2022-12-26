import os
def rand_select(option):
    import random
    rand = random.random()
    for i in range(len(option)):
        if rand < 1 / len(option) * (i + 1):
            return option[i]

def create_dir(dir_path):
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)

def word(text):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    import random

    rand = int(random.random() * 200 - 100)
    font = ["萌妹子体", "李国夫手写体", "陈静的字完整版"]
    doc = Document()
    para = doc.add_paragraph(" ")
    for line in text:
        for i in line:
            run = para.add_run(i)
            run.font.name = rand_select(font)
            run.font.size = 300000 + rand
            run.font.color.rgb = RGBColor(60, 60, 60)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
            rand = int(random.random() * 200 - 100)
    doc.save(os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.docx")

def word_to_jpg():
    from docx2pdf import convert
    
    from PyPDF2 import PdfReader
    convert(os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.docx", os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.pdf")
    reader = PdfReader(os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.pdf")
    if reader.is_encrypted:
        reader.decrypt('')
    page = len(reader.pages)
    for i in range(page):
        os.system(os.path.split(os.path.realpath(__file__))[0] + "\\mutool.exe draw -o " + os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text" + str(i + 1) + ".png -w 1998 -h 2585 " + os.path.split(os.path.realpath(__file__))[0] + "\\temporary\\text.pdf " + str(i + 1))
    return page

def pho_mix(page):
    import cv2
    import copy
    import random as ra
    background = cv2.imread(os.path.split(os.path.realpath(__file__))[0] + '\\background.jpg')
    offset_x = 30
    offset_y = 100
    for num in range(page):
        word = cv2.imread(os.path.split(os.path.realpath(__file__))[0] + '\\temporary\\text' + str(num + 1) + '.png')
        word = cv2.resize(word, (0, 0), fx = 1.1, fy = 1.2)
        res = copy.deepcopy(background)
        print('start ' + str(num + 1) + '/' + str(page))
        for i in range(len(word)):
            for j in range(len(word[0])):
                for k in range(3):
                    try:
                        res[i][j][k] = int(res[i][j][k]) * int(word[i - offset_y][j - offset_x][k]) / 255
                    except:
                        pass
        print('finish ' + str(num + 1))
        rows, cols, channel = res.shape
        M = cv2.getRotationMatrix2D((cols / 2, rows / 2), ra.random() * 8 - 4, 1.085)
        res = cv2.warpAffine(res, M, (cols, rows))
        cv2.imwrite(os.path.split(os.path.realpath(__file__))[0] + '\\res\\res' + str(num + 1) + '.jpg', res)

def main():
    
    path = os.path.split(os.path.realpath(__file__))[0]
    create_dir(path + '\\res')
    create_dir(path + '\\temporary')
    for i in os.listdir(path + '\\res'):
        os.remove(path + '\\res\\' + i)
    with open(path + "\\input.txt", "r", encoding='utf-8') as f:
        word(f.readlines())
    if input("continue?(Y/N)") == "Y":
        pass
    else:
        return
    pho_mix(word_to_jpg())
    for i in os.listdir(path + '\\temporary'):
        os.remove(os.path.split(os.path.realpath(__file__))[0] + '\\temporary\\'+ i)

if __name__=='__main__':
    main()
