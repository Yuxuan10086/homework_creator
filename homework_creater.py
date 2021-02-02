def rand_select(option):
    import random
    rand = random.random()
    for i in range(len(option)):
        if rand < 1 / len(option) * (i + 1):
            return option[i]

def word(text):
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    font = ["萌妹子体", "李国夫手写体", "陈静的字完整版"]
    doc = Document()
    para = doc.add_paragraph(" ")
    for line in text:
        for i in line:
            run = para.add_run(i)
            run.font.name = rand_select(font)
            run.font.size = 300000
            run.font.color.rgb = RGBColor(60, 60, 60)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
    doc.save("temporary\\text.docx")

with open("input.txt", "r", encoding='utf-8') as f:
     word(f.readlines())